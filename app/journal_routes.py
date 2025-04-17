from flask import Blueprint, render_template, request, jsonify, g, redirect, url_for
from app.models import JournalEntry, db
from datetime import datetime, date
from functools import wraps
import json
import io
import calendar

# Optional imports for PDF/DOCX export
# These features will be disabled if the packages aren't installed
HAS_PDF_SUPPORT = False
HAS_DOCX = False

try:
    from flask import make_response, send_file
    from weasyprint import HTML
    HAS_PDF_SUPPORT = True
except ImportError:
    # weasyprint is not installed
    pass

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    HAS_DOCX = True
except ImportError:
    # python-docx or BeautifulSoup is not installed
    pass

# Create a separate blueprint for journal routes
journal_bp = Blueprint('journal', __name__, url_prefix='/journal')

# Import the login_required decorator from routes
from app.routes import login_required

# Route to render the daily journal page
@journal_bp.route('/')
@login_required
def daily_journal():
    """Redirect to The Journal page."""
    return redirect(url_for('journal.the_journal'))

# Route to render The Journal page
@journal_bp.route('/the-journal')
@login_required
def the_journal():
    """Render The Journal page with calendar and advanced editor."""
    return render_template('journal.html')

# API route to save a journal entry
@journal_bp.route('/api/save', methods=['POST'])
@login_required
def save_journal():
    """Save a journal entry to the database."""
    if not request.is_json:
        return jsonify({'success': False, 'message': 'Invalid request format'}), 400
        
    data = request.json
    date_str = data.get('date')
    title = data.get('title', '').strip()
    content = data.get('content', '').strip()
    goals = data.get('goals', '').strip()
    mood = data.get('mood', '').strip()
    
    # Validate required fields
    if not date_str:
        return jsonify({'success': False, 'message': 'Date is required'}), 400
    if not title:
        return jsonify({'success': False, 'message': 'Title is required'}), 400
    if not content:
        return jsonify({'success': False, 'message': 'Content is required'}), 400
    if not mood:
        return jsonify({'success': False, 'message': 'Mood is required'}), 400
        
    try:
        # Parse the date string
        entry_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        # Check if an entry already exists for this date
        existing_entry = JournalEntry.query.filter_by(
            user_id=g.user.id, 
            date=entry_date
        ).first()
        
        if existing_entry:
            # Update existing entry
            existing_entry.title = title
            existing_entry.content = content
            existing_entry.goals = goals
            existing_entry.mood = mood
            existing_entry.updated_at = datetime.utcnow()
            
            db.session.commit()
            return jsonify({
                'success': True, 
                'message': 'Journal entry updated successfully',
                'entry_id': existing_entry.id
            })
        else:
            # Create new entry
            new_entry = JournalEntry(
                user_id=g.user.id,
                date=entry_date,
                entry_date=entry_date,  # For backwards compatibility
                title=title,
                content=content,
                goals=goals,
                mood=mood
            )
            
            db.session.add(new_entry)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'message': 'Journal entry created successfully',
                'entry_id': new_entry.id
            })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error saving journal entry: {str(e)}'}), 500

# API route to get all journal entries
@journal_bp.route('/api/entries')
@login_required
def get_journal_entries():
    """Get all journal entries for the current user."""
    try:
        entries = JournalEntry.query.filter_by(user_id=g.user.id).order_by(JournalEntry.date.desc()).all()
        
        entries_data = []
        for entry in entries:
            entries_data.append({
                'id': entry.id,
                'date': entry.date.strftime('%Y-%m-%d'),
                'title': entry.title,
                'content': entry.content,
                'goals': entry.goals,
                'mood': entry.mood,
                'created_at': entry.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': entry.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            })
            
        return jsonify({
            'success': True,
            'entries': entries_data
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error retrieving journal entries: {str(e)}'}), 500

# API route to get a specific journal entry by date
@journal_bp.route('/api/entry')
@login_required
def get_journal_entry():
    """Get a journal entry for a specific date."""
    date_str = request.args.get('date')
    
    if not date_str:
        return jsonify({'success': False, 'message': 'Date parameter is required'}), 400
        
    try:
        entry_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        entry = JournalEntry.query.filter_by(
            user_id=g.user.id, 
            date=entry_date
        ).first()
        
        if entry:
            entry_data = {
                'id': entry.id,
                'date': entry.date.strftime('%Y-%m-%d'),
                'title': entry.title,
                'content': entry.content,
                'goals': entry.goals,
                'mood': entry.mood,
                'created_at': entry.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': entry.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            }
            
            return jsonify({
                'success': True,
                'entry': entry_data
            })
        else:
            return jsonify({
                'success': True,
                'entry': None
            })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error retrieving journal entry: {str(e)}'}), 500

# New API routes for The Journal page

# Get a note for The Journal by date
@journal_bp.route('/api/note')
@login_required
def get_note():
    """Get a note for a specific date for The Journal."""
    date_str = request.args.get('date')
    
    if not date_str:
        return jsonify({'success': False, 'message': 'Date parameter is required'}), 400
        
    try:
        entry_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        entry = JournalEntry.query.filter_by(
            user_id=g.user.id, 
            date=entry_date
        ).first()
        
        if entry:
            entry_data = {
                'id': entry.id,
                'date': entry.date.strftime('%Y-%m-%d'),
                'title': entry.title,
                'content': entry.content
            }
            
            return jsonify({
                'success': True,
                'note': entry_data
            })
        else:
            return jsonify({
                'success': True,
                'note': None
            })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error retrieving note: {str(e)}'}), 500

# Save a note for The Journal
@journal_bp.route('/api/save-note', methods=['POST'])
@login_required
def save_note():
    """Save a note to the database for The Journal."""
    if not request.is_json:
        return jsonify({'success': False, 'message': 'Invalid request format'}), 400
        
    data = request.json
    date_str = data.get('date')
    title = data.get('title', 'Untitled Note').strip()
    content = data.get('content', '').strip()
    
    # Validate required fields
    if not date_str:
        return jsonify({'success': False, 'message': 'Date is required'}), 400
        
    try:
        # Parse the date string
        entry_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        # Check if an entry already exists for this date
        existing_entry = JournalEntry.query.filter_by(
            user_id=g.user.id, 
            date=entry_date
        ).first()
        
        if existing_entry:
            # Update existing entry
            existing_entry.title = title
            existing_entry.content = content
            existing_entry.updated_at = datetime.utcnow()
            
            db.session.commit()
            return jsonify({
                'success': True, 
                'message': 'Note updated successfully',
                'entry_id': existing_entry.id
            })
        else:
            # Create new entry
            new_entry = JournalEntry(
                user_id=g.user.id,
                date=entry_date,
                entry_date=entry_date,  # For backwards compatibility
                title=title,
                content=content,
                mood='neutral'  # Default mood for compatibility
            )
            
            db.session.add(new_entry)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'message': 'Note created successfully',
                'entry_id': new_entry.id
            })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error saving note: {str(e)}'}), 500

# Get notes for a specific month
@journal_bp.route('/api/notes-by-month')
@login_required
def get_notes_by_month():
    """Get all notes for a specific month."""
    try:
        year = int(request.args.get('year', datetime.now().year))
        month = int(request.args.get('month', datetime.now().month))
        
        # Get first and last day of the month
        first_day = date(year, month, 1)
        _, last_day_num = calendar.monthrange(year, month)
        last_day = date(year, month, last_day_num)
        
        # Query entries for the month
        entries = JournalEntry.query.filter(
            JournalEntry.user_id == g.user.id,
            JournalEntry.date >= first_day,
            JournalEntry.date <= last_day
        ).all()
        
        notes_data = []
        for entry in entries:
            notes_data.append({
                'id': entry.id,
                'date': entry.date.strftime('%Y-%m-%d'),
                'title': entry.title
            })
            
        return jsonify({
            'success': True,
            'notes': notes_data
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error retrieving notes: {str(e)}'}), 500

# Export note as PDF
@journal_bp.route('/api/export-pdf', methods=['POST'])
@login_required
def export_pdf():
    """Export a note as PDF document."""
    if not HAS_PDF_SUPPORT:
        return jsonify({'success': False, 'message': 'PDF export is not available. Please install weasyprint.'}), 501
    
    try:
        from flask import make_response
        title = request.form.get('title', 'Untitled Note')
        content = request.form.get('content', '')
        date_str = request.form.get('date', datetime.now().strftime('%Y-%m-%d'))
        
        # Create HTML from the note
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; }}
                .date {{ color: #666; margin-bottom: 20px; }}
                .content {{ line-height: 1.6; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            <div class="date">Date: {date_str}</div>
            <div class="content">{content}</div>
        </body>
        </html>
        """
        
        # Generate PDF
        pdf = HTML(string=html_content).write_pdf()
        
        # Create response
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=Journal_{date_str}_{title.replace(" ", "_")}.pdf'
        
        return response
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating PDF: {str(e)}'}), 500

# Export note as DOCX
@journal_bp.route('/api/export-docx', methods=['POST'])
@login_required
def export_docx():
    """Export a note as DOCX document."""
    if not HAS_DOCX:
        return jsonify({'success': False, 'message': 'DOCX export is not available. Please install python-docx and beautifulsoup4.'}), 501
        
    try:
        from flask import send_file
        title = request.form.get('title', 'Untitled Note')
        content = request.form.get('content', '')
        date_str = request.form.get('date', datetime.now().strftime('%Y-%m-%d'))
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Date: {date_str}")
        date_run.italic = True
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a separator
        doc.add_paragraph("").add_run("_" * 50)
        
        # Parse HTML content with BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        
        # Process the HTML elements and add to document
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
            if element.name.startswith('h'):
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.bold = True
                size = 16 - int(element.name[1])  # h1=15pt, h2=14pt, etc.
                run.font.size = Pt(max(size, 11))
            elif element.name == 'p':
                p = doc.add_paragraph()
                # Handle formatting within paragraph
                for child in element.children:
                    if child.name == 'strong' or child.name == 'b':
                        run = p.add_run(child.get_text())
                        run.bold = True
                    elif child.name == 'em' or child.name == 'i':
                        run = p.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'u':
                        run = p.add_run(child.get_text())
                        run.underline = True
                    else:
                        # Plain text or unknown tag
                        text = child.get_text() if hasattr(child, 'get_text') else str(child)
                        p.add_run(text)
            # Skip list handling for simplicity
        
        # Save to a BytesIO object
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        # Return the file
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Journal_{date_str}_{title.replace(" ", "_")}.docx'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating DOCX: {str(e)}'}), 500
